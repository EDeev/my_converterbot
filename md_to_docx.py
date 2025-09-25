#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import re
import sys
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.section import WD_SECTION
from docx.oxml.shared import OxmlElement, qn
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml


class DocumentSettings:
    """Настройки форматирования документа с поддержкой ГОСТ"""
    def __init__(self):
        # Базовые настройки текста
        self.font_name = "Times New Roman"
        self.font_size = 14  # основной текст
        self.line_spacing = 1.5
        self.justify_text = True
        self.paragraph_spacing = 6
        self.text_color = (0, 0, 0)
        self.paragraph_indent = 1.25
        
        # Отступы от полей документа в сантиметрах (ГОСТ 7.32-2017)
        self.margin_top = 2.0
        self.margin_bottom = 2.0
        self.margin_left = 3.0    # увеличено для переплета
        self.margin_right = 1.5
        
        # Настройки шрифтов заголовков по ГОСТ
        self.heading1_font_size = 16  # Заголовки глав
        self.heading2_font_size = 14  # Заголовки разделов  
        self.heading3_font_size = 14  # Подзаголовки
        self.heading4_font_size = 14
        self.heading5_font_size = 12
        self.heading6_font_size = 12
        self.footnote_font_size = 10
        
        # Интервалы заголовков по ГОСТ
        self.heading_spacing_before = 12  # пт
        self.heading_spacing_after = 6    # пт
        self.paragraph_spacing_before = 0 # пт
        
        # Нумерация страниц
        self.page_numbering = True
        self.page_number_position = "bottom_center"  # top_right, bottom_center, bottom_right
        self.page_number_start = 1
        self.exclude_title_page_numbering = True
        
        # Автонумерация заголовков
        self.auto_numbering_headings = False
        self.numbering_format = "decimal"  # "decimal" (1.1.1) или "simple" (1)
        
        # Дополнительные ГОСТ настройки
        self.bibliography_style = "gost"
        self.table_caption_position = "above"  # above, below
        self.figure_caption_position = "below"



class MarkdownToDocxConverter:
    """Конвертер Markdown в DOCX с поддержкой ГОСТ"""
    
    def __init__(self, settings: DocumentSettings = None):
        self.settings = settings or DocumentSettings()
        self.doc = Document()
        
        # Счетчики для автонумерации
        self.heading_counters = [0] * 6  # для 6 уровней заголовков
        self.footnote_counter = 0
        self.table_counter = 0
        self.figure_counter = 0
        
        self.setup_document_margins()
        self.setup_page_numbering()
        self.setup_styles()
        
    def setup_document_margins(self):
        """Настройка отступов от полей документа по ГОСТ"""
        sections = self.doc.sections
        for section in sections:
            section.top_margin = Cm(self.settings.margin_top)
            section.bottom_margin = Cm(self.settings.margin_bottom)
            section.left_margin = Cm(self.settings.margin_left)
            section.right_margin = Cm(self.settings.margin_right)
            
    def setup_page_numbering(self):
        """Настройка нумерации страниц согласно ГОСТ"""
        if not self.settings.page_numbering:
            return
            
        section = self.doc.sections[0]
        
        # Создание колонтитула для нумерации
        if self.settings.page_number_position == "bottom_center":
            footer = section.footer
            footer_para = footer.paragraphs[0]
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
        elif self.settings.page_number_position == "top_right":
            header = section.header
            header_para = header.paragraphs[0]
            header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
        elif self.settings.page_number_position == "bottom_right":
            footer = section.footer
            footer_para = footer.paragraphs[0]
            footer_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    def setup_styles(self):
        """Настройка стилей документа в соответствии с ГОСТ"""
        styles = self.doc.styles
        
        # Настройка базового стиля
        normal_style = styles['Normal']
        normal_font = normal_style.font
        normal_font.name = self.settings.font_name
        normal_font.size = Pt(self.settings.font_size)
        normal_font.color.rgb = RGBColor(*self.settings.text_color)
        
        normal_paragraph = normal_style.paragraph_format
        normal_paragraph.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        normal_paragraph.line_spacing = self.settings.line_spacing
        normal_paragraph.space_after = Pt(self.settings.paragraph_spacing)
        normal_paragraph.space_before = Pt(self.settings.paragraph_spacing_before)
        normal_paragraph.first_line_indent = Cm(self.settings.paragraph_indent)
        
        if self.settings.justify_text:
            normal_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
        # Настройка стилей заголовков с дифференцированными размерами
        heading_sizes = [
            self.settings.heading1_font_size,
            self.settings.heading2_font_size,
            self.settings.heading3_font_size,
            self.settings.heading4_font_size,
            self.settings.heading5_font_size,
            self.settings.heading6_font_size
        ]
        
        for i in range(1, 7):
            heading_style_name = f'Heading {i}'
            if heading_style_name in [s.name for s in styles]:
                heading_style = styles[heading_style_name]
            else:
                heading_style = styles.add_style(heading_style_name, WD_STYLE_TYPE.PARAGRAPH)
                
            heading_font = heading_style.font
            heading_font.name = self.settings.font_name
            heading_font.size = Pt(heading_sizes[i-1])  # используем соответствующий размер
            heading_font.bold = True
            heading_font.color.rgb = RGBColor(*self.settings.text_color)
            
            heading_paragraph = heading_style.paragraph_format
            heading_paragraph.space_before = Pt(self.settings.heading_spacing_before)
            heading_paragraph.space_after = Pt(self.settings.heading_spacing_after)
            
            # Заголовки 1 и 2 уровня по центру (ГОСТ), остальные с отступом
            if i <= 2:
                heading_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                heading_paragraph.first_line_indent = Cm(0)
            else:
                if self.settings.justify_text:
                    heading_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                heading_paragraph.first_line_indent = Cm(self.settings.paragraph_indent)
        
        # Стиль для сносок
        try:
            footnote_style = styles.add_style('Footnote', WD_STYLE_TYPE.PARAGRAPH)
            footnote_font = footnote_style.font
            footnote_font.name = self.settings.font_name
            footnote_font.size = Pt(self.settings.footnote_font_size)
            footnote_font.color.rgb = RGBColor(*self.settings.text_color)
            
            footnote_paragraph = footnote_style.paragraph_format
            footnote_paragraph.space_before = Pt(3)
            footnote_paragraph.space_after = Pt(3)
            footnote_paragraph.first_line_indent = Cm(0.5)
        except:
            pass
            
        # Стиль для кода (без изменений)
        try:
            code_style = styles.add_style('Code', WD_STYLE_TYPE.CHARACTER)
            code_font = code_style.font
            code_font.name = 'Courier New'
            code_font.size = Pt(self.settings.font_size)
            code_font.color.rgb = RGBColor(*self.settings.text_color)
        except:
            pass
            
        # Стиль для блоков кода
        try:
            code_block_style = styles.add_style('Code Block', WD_STYLE_TYPE.PARAGRAPH)
            code_block_font = code_block_style.font
            code_block_font.name = 'Courier New'
            code_block_font.size = Pt(self.settings.font_size)
            code_block_font.color.rgb = RGBColor(*self.settings.text_color)
            
            code_block_paragraph = code_block_style.paragraph_format
            code_block_paragraph.left_indent = Inches(0.5)
            code_block_paragraph.first_line_indent = Cm(0)  # без отступа первой строки для кода
            code_block_paragraph.space_before = Pt(6)
            code_block_paragraph.space_after = Pt(6)
        except:
            pass
            
        # Стиль для подписей к таблицам и рисункам
        try:
            caption_style = styles.add_style('Caption', WD_STYLE_TYPE.PARAGRAPH)
            caption_font = caption_style.font
            caption_font.name = self.settings.font_name
            caption_font.size = Pt(self.settings.font_size - 2)  # меньше основного текста
            caption_font.color.rgb = RGBColor(*self.settings.text_color)
            
            caption_paragraph = caption_style.paragraph_format
            caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption_paragraph.space_before = Pt(6)
            caption_paragraph.space_after = Pt(6)
        except:
            pass
    
    def generate_heading_number(self, level: int) -> str:
        """Генерация номера заголовка согласно настройкам автонумерации"""
        if not self.settings.auto_numbering_headings:
            return ""
        
        # Обновляем счетчик текущего уровня
        self.heading_counters[level - 1] += 1
        
        # Обнуляем счетчики всех нижестоящих уровней
        for i in range(level, 6):
            self.heading_counters[i] = 0
        
        if self.settings.numbering_format == "simple":
            return f"{self.heading_counters[level - 1]}. "
        else:  # decimal
            # Формируем иерархическую нумерацию
            numbers = []
            for i in range(level):
                if self.heading_counters[i] > 0:
                    numbers.append(str(self.heading_counters[i]))
            return ".".join(numbers) + ". " if numbers else ""
    
    def parse_markdown_file(self, file_path: str):
        """Чтение и парсинг Markdown файла"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
            return content
        except Exception as e:
            raise Exception(f"Ошибка чтения файла: {e}")
    
    def add_text_run_with_color(self, paragraph, text, bold=False, italic=False, code_style=False):
        """Добавление текста с настройкой цвета"""
        run = paragraph.add_run(text)
        run.font.color.rgb = RGBColor(*self.settings.text_color)
        
        if bold:
            run.font.bold = True
        if italic:
            run.font.italic = True
        if code_style:
            run.style = 'Code'
        
        return run
    
    def process_text_formatting(self, text: str, paragraph):
        """Обработка форматирования текста включая сноски [^1]"""
        # Обработка сносок
        footnote_pattern = r'\[\^(\d+)\]'
        footnotes = re.findall(footnote_pattern, text)
        
        # Заменяем сноски на верхние индексы
        for footnote_num in footnotes:
            text = re.sub(rf'\[\^{footnote_num}\]', f'{{FOOTNOTE_{footnote_num}}}', text)
        
        # Разбор текста на части с различным форматированием
        parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|`.*?`|\{FOOTNOTE_\d+\})', text)
        
        for part in parts:
            if not part:
                continue
            
            if part.startswith('**') and part.endswith('**'):
                # Жирный текст
                self.add_text_run_with_color(paragraph, part[2:-2], bold=True)
            elif part.startswith('*') and part.endswith('*'):
                # Курсив
                self.add_text_run_with_color(paragraph, part[1:-1], italic=True)
            elif part.startswith('`') and part.endswith('`'):
                # Инлайн код
                self.add_text_run_with_color(paragraph, part[1:-1], code_style=True)
            elif part.startswith('{FOOTNOTE_') and part.endswith('}'):
                # Сноска - добавляем как верхний индекс
                footnote_num = re.search(r'FOOTNOTE_(\d+)', part).group(1)
                run = self.add_text_run_with_color(paragraph, footnote_num)
                run.font.superscript = True
            else:
                # Обычный текст
                self.add_text_run_with_color(paragraph, part)
    
    def process_list(self, lines: list, start_idx: int):
        """Обработка списков с правильным форматированием по ГОСТ"""
        i = start_idx
        list_items = []
        
        while i < len(lines):
            line = lines[i].strip()
            
            if re.match(r'^[-*+]\s', line):
                item_text = re.sub(r'^[-*+]\s', '', line)
                list_items.append(('bullet', item_text, 0))
            elif re.match(r'^\d+\.\s', line):
                item_text = re.sub(r'^\d+\.\s', '', line)
                list_items.append(('number', item_text, 0))
            elif re.match(r'^  [-*+]\s', line):
                item_text = re.sub(r'^  [-*+]\s', '', line)
                list_items.append(('bullet', item_text, 1))
            elif re.match(r'^  \d+\.\s', line):
                item_text = re.sub(r'^  \d+\.\s', '', line)
                list_items.append(('number', item_text, 1))
            elif line == '':
                i += 1
                continue
            else:
                break
            i += 1
        
        # Добавление элементов списка с настройками ГОСТ
        for list_type, text, level in list_items:
            paragraph = self.doc.add_paragraph()
            paragraph.paragraph_format.left_indent = Cm(level * 0.75)  # увеличенный отступ для вложенности
            paragraph.paragraph_format.first_line_indent = Cm(self.settings.paragraph_indent)
            
            if list_type == 'bullet':
                paragraph.style = 'List Bullet'
                # Используем тире вместо точек (согласно ГОСТ)
                bullet_run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
                bullet_run.text = "– "  # длинное тире
            else:
                paragraph.style = 'List Number'
            
            self.process_text_formatting(text, paragraph)
        
        return i - 1
    
    def process_table(self, lines: list, start_idx: int):
        """Обработка таблиц с подписями согласно ГОСТ"""
        i = start_idx
        table_lines = []
        
        while i < len(lines):
            line = lines[i].strip()
            if '|' in line:
                table_lines.append(line)
            elif line == '':
                i += 1
                continue
            else:
                break
            i += 1
        
        if len(table_lines) < 2:
            return start_idx
        
        # Добавляем подпись к таблице (если настроено)
        if self.settings.table_caption_position == "above":
            self.table_counter += 1
            caption_para = self.doc.add_paragraph()
            caption_para.style = 'Caption'
            caption_para.add_run(f"Таблица {self.table_counter}")
        
        # Парсинг и создание таблицы
        headers = [cell.strip() for cell in table_lines[0].split('|')[1:-1]]
        data_lines = table_lines[2:] if len(table_lines) > 2 else []
        
        table = self.doc.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'
        
        # Заполнение заголовков
        header_row = table.rows[0]
        for idx, header in enumerate(headers):
            cell = header_row.cells[idx]
            cell.text = header
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.bold = True
        
        # Заполнение данных
        for line in data_lines:
            row_data = [cell.strip() for cell in line.split('|')[1:-1]]
            row = table.add_row()
            for idx, cell_data in enumerate(row_data):
                if idx < len(row.cells):
                    row.cells[idx].text = cell_data
                    # Выравнивание по центру для всех ячеек (ГОСТ)
                    for paragraph in row.cells[idx].paragraphs:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Подпись снизу (если настроено)
        if self.settings.table_caption_position == "below":
            self.table_counter += 1
            caption_para = self.doc.add_paragraph()
            caption_para.style = 'Caption'
            caption_para.add_run(f"Таблица {self.table_counter}")
        
        return i - 1
    
    def process_code_block(self, lines: list, start_idx: int):
        """Обработка блоков кода"""
        i = start_idx + 1
        code_lines = []
        
        while i < len(lines):
            line = lines[i]
            if line.strip().startswith('```'):
                break
            code_lines.append(line)
            i += 1
        
        code_paragraph = self.doc.add_paragraph()
        code_paragraph.style = 'Code Block'
        code_paragraph.add_run('\n'.join(code_lines))
        
        return i
    
    def add_footnote_definition(self, footnote_num: str, footnote_text: str):
        """Добавление определения сноски в конец документа"""
        footnote_para = self.doc.add_paragraph()
        footnote_para.style = 'Footnote'
        
        # Номер сноски как верхний индекс
        footnote_run = footnote_para.add_run(footnote_num)
        footnote_run.font.superscript = True
        
        # Текст сноски
        footnote_para.add_run(f" {footnote_text}")
    
    def process_bibliography(self, lines: list, start_idx: int):
        """Обработка списка литературы в стиле ГОСТ"""
        i = start_idx
        bib_items = []
        
        # Поиск элементов библиографии
        while i < len(lines):
            line = lines[i].strip()
            if re.match(r'^\d+\.\s', line):
                bib_text = re.sub(r'^\d+\.\s', '', line)
                bib_items.append(bib_text)
            elif line == '':
                i += 1
                continue
            else:
                break
            i += 1
        
        if bib_items:
            # Заголовок списка литературы
            bib_heading = self.doc.add_paragraph()
            bib_heading.style = 'Heading 1'
            bib_heading.add_run("СПИСОК ЛИТЕРАТУРЫ")
            
            # Элементы библиографии
            for idx, item in enumerate(bib_items, 1):
                bib_para = self.doc.add_paragraph()
                bib_para.paragraph_format.first_line_indent = Cm(0)
                bib_para.paragraph_format.left_indent = Cm(1)
                bib_para.add_run(f"{idx}. {item}")
        
        return i - 1
    
    def convert(self, md_file_path: str, output_path: str = None):
        """Основной метод конвертации с поддержкой ГОСТ"""
        if not output_path:
            md_path = Path(md_file_path)
            output_path = md_path.with_suffix('.docx')
        
        content = self.parse_markdown_file(md_file_path)
        lines = content.split('\n')
        
        # Сбор сносок для обработки в конце
        footnote_definitions = {}
        
        i = 0
        while i < len(lines):
            line = lines[i]
            stripped_line = line.strip()
            
            if not stripped_line:
                i += 1
                continue
            
            # Обработка определений сносок [^1]: текст сноски
            footnote_def_match = re.match(r'^\[\^(\d+)\]:\s*(.+)', stripped_line)
            if footnote_def_match:
                footnote_num = footnote_def_match.group(1)
                footnote_text = footnote_def_match.group(2)
                footnote_definitions[footnote_num] = footnote_text
                i += 1
                continue
            
            # Заголовки с автонумерацией
            if stripped_line.startswith('#'):
                match = re.match(r'^(#{1,6})\s+(.+)', stripped_line)
                if match:
                    level = len(match.group(1))
                    title = match.group(2)
                    
                    # Разрыв страницы перед заголовком 2 уровня
                    if level == 2:
                        self.doc.add_page_break()
                    
                    heading = self.doc.add_paragraph()
                    heading.style = f'Heading {level}'
                    
                    # Добавляем автонумерацию
                    heading_number = self.generate_heading_number(level)
                    full_title = heading_number + title
                    
                    self.process_text_formatting(full_title, heading)
            
            # Блоки кода
            elif stripped_line.startswith('```'):
                i = self.process_code_block(lines, i)
            
            # Таблицы
            elif '|' in stripped_line:
                i = self.process_table(lines, i)
            
            # Списки
            elif re.match(r'^[-*+]\s', stripped_line) or re.match(r'^\d+\.\s', stripped_line):
                i = self.process_list(lines, i)
            
            # Список литературы (если заголовок содержит "литература" или "bibliography")
            elif re.match(r'^#+\s*(список\s+литературы|bibliography|references)', stripped_line, re.IGNORECASE):
                i = self.process_bibliography(lines, i + 1)
            
            # Цитаты
            elif stripped_line.startswith('>'):
                quote_text = re.sub(r'^>\s?', '', stripped_line)
                quote_paragraph = self.doc.add_paragraph()
                quote_paragraph.paragraph_format.left_indent = Inches(0.5)
                quote_paragraph.paragraph_format.right_indent = Inches(0.5)
                self.process_text_formatting(quote_text, quote_paragraph)
                
                for run in quote_paragraph.runs:
                    run.font.italic = True
            
            # Горизонтальные линии
            elif stripped_line in ['---', '***', '___']:
                hr_paragraph = self.doc.add_paragraph()
                hr_paragraph.add_run('_' * 50)
                hr_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Обычные абзацы
            else:
                paragraph = self.doc.add_paragraph()
                self.process_text_formatting(stripped_line, paragraph)
            
            i += 1
        
        # Добавление сносок в конец документа
        if footnote_definitions:
            # Разделительная линия
            self.doc.add_paragraph().add_run('_' * 50)
            
            for footnote_num in sorted(footnote_definitions.keys(), key=int):
                self.add_footnote_definition(footnote_num, footnote_definitions[footnote_num])
        
        self.doc.save(output_path)
        return output_path


def main():
    """Основная функция для запуска из командной строки"""
    if len(sys.argv) < 2:
        print("Использование: python md_converter.py <путь_к_md_файлу> [путь_к_выходному_файлу]")
        return
    
    md_file = sys.argv[1]
    output_file = sys.argv[2] if len(sys.argv) > 2 else None
    
    # ГОСТ-совместимые настройки по умолчанию
    settings = DocumentSettings()
    
    converter = MarkdownToDocxConverter(settings)
    
    try:
        output_path = converter.convert(md_file, output_file)
        print(f"Файл успешно конвертирован: {output_path}")
    except Exception as e:
        print(f"Ошибка конвертации: {e}")


if __name__ == "__main__":
    main()


# Пример использования с кастомными ГОСТ настройками:
"""
settings = DocumentSettings()
settings.font_name = "Times New Roman"
settings.font_size = 14
settings.heading1_font_size = 16
settings.heading2_font_size = 14
settings.line_spacing = 1.5
settings.margin_left = 3.0  # для переплета
settings.auto_numbering_headings = True
settings.numbering_format = "decimal"  # 1.1.1 формат
settings.page_numbering = True
settings.page_number_position = "bottom_center"

converter = MarkdownToDocxConverter(settings)
converter.convert("dissertation.md", "dissertation_gost.docx")
"""